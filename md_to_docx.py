"""
Convierte Plan_Sucesiones_Unidad3_Integracion.md a .docx con estilos de Word reales.
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_FILE  = "Plan_Sucesiones_Unidad3_Integracion.md"
OUTPUT_FILE = "Plan_Sucesiones_Unidad3_Integracion.docx"

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=None, color=None, mono=False):
    run.bold = bold
    run.italic = italic
    if mono:
        run.font.name = "Courier New"
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def shade_paragraph(paragraph, fill_hex="F2F2F2"):
    """Aplica fondo gris claro al párrafo (para bloques de código)."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)

def add_border_bottom(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4472C4")
    pBdr.append(bottom)
    pPr.append(pBdr)

def apply_inline(paragraph, text):
    """Aplica negritas e itálicas inline al texto de un párrafo."""
    # Patterns: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        # texto antes del match
        if m.start() > last:
            paragraph.add_run(text[last:m.start()])
        raw = m.group(0)
        if raw.startswith("**"):
            r = paragraph.add_run(m.group(2))
            r.bold = True
        elif raw.startswith("*"):
            r = paragraph.add_run(m.group(3))
            r.italic = True
        elif raw.startswith("`"):
            r = paragraph.add_run(m.group(4))
            r.font.name = "Courier New"
            r.font.size = Pt(10)
        last = m.end()
    if last < len(text):
        paragraph.add_run(text[last:])

def parse_table(doc, lines, start):
    """Detecta y agrega una tabla markdown."""
    # Recoger todas las filas hasta que deje de ser tabla
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        # Fila separadora (|:---|:---|)
        if re.match(r'^\|[\s\-:|]+\|', line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1

    if not rows:
        return start

    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row_cells in enumerate(rows):
        for c_idx, cell_text in enumerate(row_cells):
            if c_idx >= col_count:
                break
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            # Primera fila = encabezado
            apply_inline(p, cell_text)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True

    doc.add_paragraph()  # espacio tras tabla
    return i

# ── Conversión principal ──────────────────────────────────────────────────────

def convert(input_path, output_path):
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    with open(input_path, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        raw = lines[i].rstrip("\n").rstrip("\r")
        stripped = raw.strip()

        # ── Bloque de código ─────────────────────────────────────────────────
        if stripped.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # Cerrar bloque
                for cl in code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent  = Cm(0.5)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after  = Pt(0)
                    shade_paragraph(p, "EFEFEF")
                    run = p.add_run(cl)
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                doc.add_paragraph()
                in_code_block = False
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(raw)
            i += 1
            continue

        # ── Línea horizontal ─────────────────────────────────────────────────
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            p = doc.add_paragraph()
            add_border_bottom(p)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            i += 1
            continue

        # ── Tabla ────────────────────────────────────────────────────────────
        if stripped.startswith("|"):
            i = parse_table(doc, lines, i)
            continue

        # ── Títulos ──────────────────────────────────────────────────────────
        m = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            text  = re.sub(r'[*_`]', '', m.group(2))  # limpiar markdown inline
            heading_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
            p = doc.add_heading(text, level=level)
            i += 1
            continue

        # ── Blockquote ───────────────────────────────────────────────────────
        if stripped.startswith(">"):
            text = re.sub(r'^>\s*', '', stripped)
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            p.paragraph_format.left_indent  = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            shade_paragraph(p, "E8F0FE")
            apply_inline(p, text)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # ── Lista (- o * o número) ───────────────────────────────────────────
        m_ul = re.match(r'^(\s*)([-*])\s+(.*)', raw)
        m_ol = re.match(r'^(\s*)(\d+)\.\s+(.*)', raw)
        if m_ul or m_ol:
            indent_spaces = len((m_ul or m_ol).group(1))
            text = (m_ul or m_ol).group(3) if m_ul else m_ol.group(3)
            p = doc.add_paragraph(style="List Bullet" if m_ul else "List Number")
            p.paragraph_format.left_indent   = Cm(0.5 + indent_spaces * 0.15)
            p.paragraph_format.space_before  = Pt(1)
            p.paragraph_format.space_after   = Pt(1)
            apply_inline(p, text)
            i += 1
            continue

        # ── Línea vacía ──────────────────────────────────────────────────────
        if stripped == "":
            # Solo agrega espacio si el párrafo anterior no es ya vacío
            if doc.paragraphs and doc.paragraphs[-1].text.strip() != "":
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # ── Párrafo normal ───────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        apply_inline(p, stripped)
        i += 1

    doc.save(output_path)
    print(f"✅ Archivo guardado: {output_path}")

if __name__ == "__main__":
    convert(INPUT_FILE, OUTPUT_FILE)
